# gdrive_utils.py
from googleapiclient.discovery import build
from get_credentials import get_credentials

from typing import List
from langchain_core.documents import Document
import os
import pdfplumber
import docx
import io

def get_drive_service():
    creds = get_credentials()
    service = build('drive', 'v3', credentials=creds)
    return service

class DriveLoader:
    """Classe pour charger les documents depuis Google Drive."""
    
    def __init__(self, doc_name=None):
        """Initialise le loader avec le service Google Drive."""
        self.service = get_drive_service()
        self.folder_or_doc_id = os.getenv('GOOGLE_DRIVE_DOC_ID')
        self.doc_name = doc_name

    def find_first_doc(self, folder_id: str) -> str:
        """Trouve le premier document Google Doc dans un dossier.
        
        Args:
            folder_id: ID du dossier Google Drive
            
        Returns:
            ID du premier document Google Doc trouvé
        """
        try:
            # Rechercher les fichiers dans le dossier
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and mimeType='application/vnd.google-apps.document'",
                pageSize=1,
                fields="files(id, name)"
            ).execute()
            
            files = results.get('files', [])
            if not files:
                raise ValueError(f"Aucun document Google Doc trouvé dans le dossier {folder_id}")
                
            return files[0]['id']
            
        except Exception as e:
            print(f"Erreur lors de la recherche du document: {str(e)}")
            return None
    
    def find_doc_by_name(self, folder_id: str, doc_name: str) -> str:
        """Trouve un document Google Doc par nom dans un dossier."""
        try:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and mimeType='application/vnd.google-apps.document' and name='{doc_name}'",
                pageSize=1,
                fields="files(id, name)"
            ).execute()
            files = results.get('files', [])
            if not files:
                raise ValueError(f"Aucun document nommé '{doc_name}' trouvé dans le dossier {folder_id}")
            return files[0]['id']
        except Exception as e:
            print(f"Erreur lors de la recherche du document: {str(e)}")
            return None

    def find_file_by_name(self, folder_id: str, file_name: str) -> dict:
        """Trouve le fichier nommé 'info_pour_chatbot' dans le dossier."""
        try:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and name='{file_name}'",
                pageSize=1,
                fields="files(id, name, mimeType)"
            ).execute()
            files = results.get('files', [])
            if not files:
                raise ValueError(f"Aucun fichier nommé '{file_name}' trouvé dans le dossier {folder_id}")
            return files[0]
        except Exception as e:
            print(f"Erreur lors de la recherche du fichier: {str(e)}")
            return None

    def load(self) -> List[Document]:
        """Charge le contenu du document depuis Google Drive.
        
        Returns:
            Liste de documents Langchain
        """
        try:
            file = self.service.files().get(
                fileId=self.folder_or_doc_id,   # <-- correction ici
                fields='id, name, mimeType'
            ).execute()

            # Si c'est un dossier, chercher le fichier 'info_pour_chatbot'
            if file['mimeType'] == 'application/vnd.google-apps.folder':
                print(f"L'ID {self.folder_or_doc_id} est un dossier, recherche du fichier 'info_pour_chatbot'...")
                file = self.find_file_by_name(self.folder_or_doc_id, "info_pour_chatbot")
                if not file:
                    raise ValueError("Aucun fichier 'info_pour_chatbot' trouvé dans le dossier")
            
            # Extraction selon le type
            text = ""
            if file['mimeType'] == 'application/vnd.google-apps.document':
                content = self.service.files().export(
                    fileId=file['id'],
                    mimeType='text/plain'
                ).execute()
                text = content.decode('utf-8')
            elif file['mimeType'] == 'application/pdf':
                print("PDF détecté, extraction du texte...")
                request = self.service.files().get_media(fileId=file['id'])
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                fh.seek(0)
                with pdfplumber.open(fh) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            elif file['mimeType'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                print("Word détecté, extraction du texte...")
                request = self.service.files().get_media(fileId=file['id'])
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                fh.seek(0)
                doc = docx.Document(fh)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                raise ValueError(f"Format non supporté: {file['mimeType']}")

            return [Document(
                page_content=text,
                metadata={
                    'source': f"Google Drive - {file['name']}",
                    'file_id': file['id']
                }
            )]
        except Exception as e:
            print(f"Erreur lors du chargement du fichier: {str(e)}")
            return []
